# Python 3.7
# -*- coding: utf_8 -*-
from selenium import webdriver
import time
from datetime import datetime, timedelta
import sqlite3
import docx
import os


def create_bd_file():
    """Создаем базу для хранения данных пользователя"""
    name_db = 'episodes.sqlite'
    cur_dir = os.getcwd()
    path_db = os.path.join(cur_dir, name_db)

    # проверка на существование базы данных
    if not os.path.exists(path_db):
        try:
            # создание базы данных + создание таблиц
            conn = sqlite3.connect(path_db)
            conn.executescript("PRAGMA foreign_keys = ON")
            cursor = conn.cursor()

            # создание таблиц пользователя
            cursor.executescript('BEGIN TRANSACTION;\n'

                                 'CREATE TABLE "update_time" ('
                                 '                `id` INTEGER PRIMARY KEY AUTOINCREMENT,\n'
                                 '                `last_update_time` DATETIME\n'
                                 ');\n'

                                 'CREATE TABLE "user" (\n'
                                 '                `id` INTEGER PRIMARY KEY AUTOINCREMENT,\n'
                                 '                `user_id` INTEGER(11) NOT NULL,\n'
                                 '                `login` VARCHAR(100),\n'
                                 '                `password` VARCHAR(30),\n '
                                 '                  UNIQUE(user_id, login)\n'
                                 '                  ON CONFLICT IGNORE'
                                 ');\n'

                                 'CREATE TABLE "gamers" ('
                                 '                `id` INTEGER PRIMARY KEY AUTOINCREMENT,\n'
                                 '                `user_id` INTEGER(11) NOT NULL,\n'
                                 '                `user_name` VARCHAR(200),\n'
                                 '                  UNIQUE(user_id, user_name)\n'
                                 '                  ON CONFLICT IGNORE'
                                 ');\n'

                                 'CREATE TABLE "episodes" ('
                                 '                `id` INTEGER PRIMARY KEY AUTOINCREMENT,\n'
                                 '                `episode_link` VARCHAR(200),\n'
                                 '                `episode_name` VARCHAR(8000),\n'
                                 '                `posts` INTEGER(11),\n'
                                 '                `last_post_date` DATETIME,\n'
                                 '                  UNIQUE(episode_link, episode_name)\n'
                                 '                  ON CONFLICT REPLACE'
                                 ');\n'

                                 'CREATE TABLE "posts" ('
                                 '                `id` INTEGER PRIMARY KEY AUTOINCREMENT,\n'
                                 '                `author_id` INTEGER(11) NOT NULL REFERENCES gamers(user_id),\n'
                                 '                `author_name` VARCHAR(200) REFERENCES gamers(user_name),\n'
                                 '                `episode_name` VARCHAR(8000) REFERENCES episodes(episode_name),\n'
                                 '                `post_link` VARCHAR(200),\n'
                                 '                `post_id` INTEGER(11),\n'
                                 '                `post_text` VARCHAR(65000),\n'
                                 '                `post_date` DATETIME, \n'
                                 '                  UNIQUE(post_id)\n'
                                 '                  ON CONFLICT IGNORE'
                                 ');\n'

                                 'COMMIT;\n')

            # фиксирую коммит
            conn.commit()

        except sqlite3.Error as e:
            print('Ошибка БД: ' + str(e))


def update_bd(sql_request, sql_params):
    """Универсальная функция для проброса запроса к бд"""
    name_db = 'episodes.sqlite'
    cur_dir = os.getcwd()
    path_db = os.path.join(cur_dir, name_db)

    # проверка на существование базы данных
    if os.path.exists(path_db):
        try:
            # коннект к существующей бд
            conn = sqlite3.connect(path_db, detect_types=sqlite3.PARSE_DECLTYPES)
            cursor = conn.cursor()
            # отправка запроса
            cursor.execute(sql_request, sql_params)
            # фиксирую коммит
            conn.commit()

            return cursor

        except sqlite3.Error as e:
            print('Ошибка БД: ' + str(e))


def updating_refresh_time():
    """Функция последнего обновления статистики"""
    ts = datetime.now().timestamp()
    ts = str(ts).split('.')[0]
    # запрос
    sql_request = '''INSERT INTO update_time VALUES(null, ?)'''
    # аргумент запроса
    sql_params = ([int(ts)])
    # передаем в бд
    update_bd(sql_request, sql_params)

    return True


class EpChecker(object):
    def __init__(self, *args, **kwargs):
        """Объявляем параметры вызова вебдрайвера"""
        # ua = dict(DesiredCapabilities.CHROME)
        self.options = webdriver.ChromeOptions()
        # options.add_argument('headless')
        # options.add_argument('window-size=1920x935')
        self.options.add_argument('--blink-settings=imagesEnabled=false')
        # self.options.add_experimental_option("detach", True)
        self.executable_path = './chromedriver/chromedriver.exe'

        self.driver = None

        # пользовательские данные (логин, пароль, ссылка на эпизод)
        self.user_login = None
        self.user_password = None
        self.user_id = None
        self.episode_url = None
        self.episode_name = None
        self.planned_gamers = None

    def get_user_data(self, login, password):
        """Получаем данные пользователя"""
        self.user_login = login
        self.user_password = password

        if self.user_login is None or self.user_password is None:
            print('Пользователь не передал данные для входа')
            return False

        return True

    def get_episode_url(self, url):
        """Получаем ссылку на эпизод"""
        self.episode_url = url.split('&p')[0].split('#p')[0] + '&p=-1'

        if self.episode_url is None:
            print('Ссылка на эпизод отсутствует')
            return False

        return True

    def get_planned_gamers(self, gamers):
        """Получаем ссылку на эпизод"""
        self.planned_gamers = gamers

        if self.planned_gamers is None:
            print('Игроки не заданы')
            return False

        return True

    def create_driver(self):
        """Инициализируем драйвер"""
        self.driver = webdriver.Chrome(options=self.options, executable_path=self.executable_path)

        return self.driver

    def login(self):
        """Логинимся на форум"""
        self.create_login_form()

        return True

    def go_to_url(self):
        """Переходим на страницу эпизода"""
        # time.sleep(3)
        self.driver.get(self.episode_url)

        return True

    def create_login_form(self):
        """Входим на форум при помощи формы"""
        js_code = f'''
        let form = '<form id="login" method="post" action="/login.php?action=in">\
        <input type=\"hidden\" name=\"form_sent\" value="1" \>\
        <input type=\"hidden" name="redirect_url" value="" \>\
        <input type=\"text" name="req_username" maxlength="25" value="{self.user_login}"\>\
        <input type=\"password" name="req_password" maxlength="16" value="{self.user_password}"\>\
        <input type=\"submit\" class=\"button\" name=\"login\"/>\
        </form>';

        $("#navlogin").after(form);
        $("#login input[type='submit']").click();
        '''
        self.driver.execute_script(js_code)

        return True

    def get_all_posts_count(self):
        """Получаем количество постов в эпизоде"""
        time.sleep(3)
        theme_id = self.episode_url.split('id=')[1].split('&p')[0].split('#p')[0]
        # вылавливаем из темы число сообщений
        all_posts_count = self.driver.execute_script(
            '''return (function () {let posts_number = $("#topic_t''' + theme_id +
            '''> h2 > span.item2").text().split("из ")[1]; return posts_number;}());''')
        all_posts_count = int(all_posts_count) - 1  # удаляем первопост с описанием темы
        print(all_posts_count)
        return all_posts_count

    def get_ep_name(self):
        """Получаем название эпизода"""
        # название эпизода
        ep_name = self.driver.find_element_by_css_selector('#pun-main > h1 > span').get_attribute('innerHTML')
        print(ep_name)
        # присваиваем переменной имя эпизода
        self.episode_name = ep_name

        if self.episode_name is None:
            print('Не передано имя эпизода')
            return False

        return self.episode_name

    def get_all_posts_from_ep(self):
        js_code = '''
        return (function () {
            const allGamersPosts = {};
            $(".topic .post:not(.topicpost)").html(function () {
                    let author_name = $(this).find('.pa-author > a').html();
                    let author_id = $(this).attr('data-user-id');
                    let post_id = $(this).attr('id').split('p')[1];
                    let post_date = $(this).attr('data-posted');
                    let post_text = $(this).find('.post-content > p:not(.post-sig):not(.lastedit)').text()
                    let post_link = $(this).find('h3 > span > a.permalink').attr('href');
                    
                    allGamersPosts[post_id] = [author_name, author_id, post_date, post_text, post_link];
                    });
            return allGamersPosts;
            }());
        '''
        posts_and_info = self.driver.execute_script(js_code)
        # print(posts_and_info)

        self.save_posts_in_bd(posts_and_info)

    def save_posts_in_bd(self, post_massive):
        """Пишем данные эпизода в таблицу эпизодов в БД"""
        for key in post_massive:
            # запрос
            sql_request = '''INSERT OR REPLACE INTO posts VALUES(null, ?, ?, ?, ?, ?, ?, ?)'''
            # аргумент запроса
            sql_params = ([post_massive[key][1], post_massive[key][0], self.episode_name, post_massive[key][4], key,
                           post_massive[key][3], post_massive[key][2]])
            # передаем в бд
            update_bd(sql_request, sql_params)

        return True

    def get_information_ep(self):
        """Получаем параметры эпизода,
        как то: число игроков и кол-во их постов в эпизоде, а также userID"""
        js_code = '''
            return (function () {
                let datas = [];
                const allGamersInEp = [];

                $(".topic .post:not(.topicpost) .container > .post-author > ul > li.pa-author > a").html(function () {
                    let hrf = $(this).html();
                    allGamersInEp.push(hrf);
                });

                let unicGamers = new Set(allGamersInEp);
                //let gamersArr = Array.from(unicGamers)
                //datas.push(gamersArr);

                let allPostsForUser = allGamersInEp.reduce((acc, el) => {
                acc[el] = (acc[el] || 0) + 1; return acc;
                }, {});
                datas.push(allPostsForUser);

                return datas;
                 }());
            '''
        users_and_their_posts = self.driver.execute_script(js_code)
        users = []
        for key in users_and_their_posts[0]:
            users.append(key)

        js_script = '''
            return (function () {
            const allGamersID = {};
            $(".topic .post:not(.topicpost) .container > .post-links > ul > li.pl-email.pm").html(function () {
                let hrf = $(this).html();
                let id = hrf.split('uid=')[1].split('" rel=')[0];
                let name = hrf.split('class="acchide">&nbsp;')[1].split('</span></a>')[0];
                allGamersID[name] = id;
                    });
                    
            return allGamersID;
            }());
            '''
        users_and_their_id = self.driver.execute_script(js_script)
        print(users_and_their_id)
        self.save_gamers_params(users_and_their_id)
        # считаем количество ключей
        self.check_gamers_count(users_and_their_id)
        # names_and_id = {}
        # for key in users_and_their_id:
        #     names_and_id.update({key: users_and_their_id[key]})
        # print(names_and_id)

        return True

    def check_gamers_count(self, gamers):
        """Получаем количество игроков в массиве"""
        count = len(gamers)
        print(count)
        self.count_versification(count)
        return count

    def count_versification(self, count):
        """Сравниваем кол-во игроков в массиве и кол-во заданных игроков"""
        if len(self.planned_gamers) is not count:
            print('Количество игроков в эпизоде на текущий момент не соответствует запланированному')
            return False
        elif len(self.planned_gamers) == count:
            print('Количество игроков в эпизоде соответствует запланированному')
            return True

    def save_gamers_params(self, gamers):
        """Пишем данные игроков в таблицу БД"""
        time.sleep(3)
        for key in gamers:
            # запрос
            sql_request = '''INSERT OR IGNORE INTO gamers VALUES(null, ?, ?)'''
            # аргумент запроса
            sql_params = ([gamers[key], key])
            # передаем в бд
            update_bd(sql_request, sql_params)

        return True

    def get_my_profile_id(self):
        """Получаем идентификатор профиля юзера"""
        time.sleep(3)
        js_code = '''return (function () { let myID = $('#navprofile a').attr('href'); return myID; }());'''
        profile_url = self.driver.execute_script(js_code)
        my_id = profile_url.split('id=')[1]
        print(my_id)
        self.user_id = my_id

        if self.user_id is None:
            print('Не передан userID')
            return False

        return self.user_id

    def save_user_params(self):
        """Пишем данные юзера в таблицу юзеров БД"""
        # запрос
        sql_request = '''INSERT OR IGNORE INTO user VALUES(null, ?, ?, ?)'''
        # аргумент запроса
        sql_params = ([self.user_id, self.user_login, self.user_password])
        # передаем в бд
        update_bd(sql_request, sql_params)

        return True

    def get_episode_params(self):
        """Переходим в последний пост и получаем параметры эпизода, как то: количество постов, дата последнего поста"""
        url = self.episode_url.split('&p')[0] + '&action=last'
        self.driver.get(url)
        js_code = '''return (function () { 
        let lastTime = $('.endpost').attr('data-posted'); 
        return lastTime; 
        }())'''
        # время последнего поста
        last_post_time = self.driver.execute_script(js_code)
        # количество постов
        posts = self.get_all_posts_count()
        # форматированная дата для информации
        print(datetime.fromtimestamp(int(last_post_time)))
        # пишем данные в бд
        self.save_ep_params(last_post_time, posts)

        return True

    def save_ep_params(self, last_post_time, posts):
        """Пишем данные эпизода в таблицу эпизодов в БД"""
        # запрос
        sql_request = '''INSERT OR REPLACE INTO episodes VALUES(null, ?, ?, ?, ?)'''
        # обрезаем url-адрес
        ep_url = self.episode_url.split('&p')[0]
        # аргумент запроса
        sql_params = ([ep_url, self.episode_name, posts, last_post_time])
        # передаем в бд
        update_bd(sql_request, sql_params)

        return True


class Base_checker(object):
    def __init__(self, ep_name, my_name, priority, strict=False):
        # входные данные по эпизоду
        self.episode_name = ep_name
        self.my_name = my_name
        self.episode_priority = priority
        self.episode_strict = strict

        # данные, получаемые в процессе обработки информации из бд
        self.last_post_author = None
        self.last_post_priority = None
        self.time_duty = None

    def post_count_in_ep(self):
        """Запрашиваем количество постов в эпизоде из таблицы постов в БД"""
        # запрос
        sql_request = '''SELECT COUNT(episode_name) from posts GROUP BY episode_name HAVING episode_name=?'''
        # аргумент запроса
        sql_params = ([self.episode_name])
        # передаем в бд
        selected = update_bd(sql_request, sql_params)

        rows = selected.fetchall()
        # если список rows не пуст, то
        if rows:
            for row in rows:
                print(f'Всего постов в эпизоде: {row[0]}')

                return row[0]
        else:
            print('Нет результата. В базе данных пусто: проверьте введенные значения')

    def last_post_date_in_ep(self):
        """Запрашиваем последний пост в эпизоде и его автора из таблицы постов в БД"""
        # запрос
        sql_request = '''SELECT author_name, post_date FROM posts WHERE episode_name=? ORDER BY post_date  DESC LIMIT 1'''
        # аргумент запроса
        sql_params = ([self.episode_name])
        # передаем в бд
        selected = update_bd(sql_request, sql_params)

        rows = selected.fetchall()
        # если список rows не пуст, то
        if rows:
            for row in rows:
                print(f'Последним написал {row[0]} в', datetime.fromtimestamp(row[1]))
                # получаем имя последнего написавшего и сохраняем его
                self.last_post_author = row[0]

                self.time_from_last_post(row[1])
                return row[0], datetime.fromtimestamp(row[1])
        else:
            print('Нет результата. В базе данных пусто: проверьте введенные значения')

    def time_from_last_post(self, last_post_date):
        """Получаем время, прошедшее с даты последнего поста"""
        this_date = datetime.now().timestamp()
        this_date = int(str(this_date).split('.')[0])
        this_date = datetime.fromtimestamp(this_date)
        delta = this_date - datetime.fromtimestamp(last_post_date)
        mm, ss = divmod(delta.seconds, 60)
        hh, mm = divmod(mm, 60)
        time_passed = 'С момента последнего поста в эпизод прошло: {} д. {} час. {} мин. {} сек.'.format(delta.days, hh,
                                                                                                         mm, ss)

        self.time_duty = time_passed
        print(time_passed)

    def my_last_post_date_in_ep(self):
        """Запрашиваем мой последний пост в эпизоде из таблицы постов в БД"""
        # запрос
        sql_request = '''SELECT author_name, post_date FROM posts WHERE episode_name=? AND author_name=? ORDER BY post_date  DESC LIMIT 1'''
        # аргумент запроса
        sql_params = ([self.episode_name, self.my_name])
        # передаем в бд
        selected = update_bd(sql_request, sql_params)

        rows = selected.fetchall()
        # если список rows не пуст, то
        if rows:
            for row in rows:
                print(f'{row[0]} последний раз написал пост в', datetime.fromtimestamp(row[1]))

                return row[0], datetime.fromtimestamp(row[1])
        else:
            print('Нет результата. В базе данных пусто: проверьте введенные значения')

    def check_strict_priority(self):
        """Проверка строгой очередности"""
        if self.episode_strict:
            print('Очередность эпизода задана как строгая, можно проверить очередь постов')
            return True
        else:
            print('Очередность эпизода задана как нестрогая, проверка невозможна')
            return False

    def get_priority(self):
        """Получаем очередность постов на текущий момент"""
        if self.episode_strict:
            # запрос
            sql_request = '''SELECT author_name FROM posts WHERE episode_name=? ORDER BY post_date DESC LIMIT ?'''
            # аргумент запроса
            sql_params = ([self.episode_name, len(self.episode_priority)])
            # передаем в бд
            selected = update_bd(sql_request, sql_params)

            rows = selected.fetchall()
            # если список rows не пуст, то
            gamers = []
            if rows:
                for row in rows:
                    gamers.append(row[0])

                gamers.reverse()
                print(gamers)

                self.last_post_priority = gamers
                return gamers
            else:
                print('Нет результата. В базе данных пусто: проверьте введенные значения')

    def check_last_post(self):
        """Проверяем, чей последний пост"""
        if self.last_post_author != self.my_name:
            # если последний пост не мой
            # проверяем, я ли должен
            self.check_debtor()
            return True
        else:
            print('Последний пост мой, проверка не требуется')
            return False

    def check_debtor(self):
        """Проверяем, моя ли очередь"""
        post_debtor = self.get_name_post_debtor()
        if post_debtor == self.my_name:
            print('Следующий в очереди мой пост')
            # если последний пост мой
            # смотрим, сколько именно я должен
            self.time_of_duty()
            return True
        else:
            print(f'Следующий в очереди не я, а {post_debtor}')
            return False

    def get_name_post_debtor(self):
        """Получаем имя следующего в очереди"""
        index_post_debtor = self.get_index_post_debtor()
        try:
            name = self.episode_priority[index_post_debtor]
        except IndexError:
            name = self.episode_priority[0]

        return name

    def get_index_post_debtor(self):
        """Получаем индекс следующего в очереди"""
        strict_index = self.search_gamer_index_in_strict()
        index_post_debtor = strict_index + 1
        return index_post_debtor

    def search_gamer_index_in_strict(self):
        """Ищем индекс последнего написавшего в очереди"""
        strict_index = self.episode_priority.index(self.last_post_author)
        return strict_index

    def time_of_duty(self):
        """Получаем время моего долга"""
        time_passed = str(self.time_duty).split('прошло: ')[1]
        print(f'Я должен уже {time_passed}')

    def get_posts_from_ep(self):
        """Запрашиваем все посты в эпизоде из таблицы постов в БД"""
        # запрос
        sql_request = '''SELECT author_name, post_date, post_text FROM posts WHERE episode_name=? ORDER BY post_date'''
        # аргумент запроса
        sql_params = ([self.episode_name])
        # передаем в бд
        selected = update_bd(sql_request, sql_params)
        # получаем выходные данные
        rows = selected.fetchall()
        # пишем вордовский файл
        self.create_word_file(rows)

    def create_word_file(self, rows):
        """Создаем вордовский файл и чистим его"""
        doc = docx.Document()
        doc.add_heading(self.episode_name, 0)
        # если список rows не пуст, то
        if rows:
            for row in rows:
                doc.add_heading(f'{row[0]}, {datetime.fromtimestamp(row[1])}', 4)
                # добавляем параграф
                doc.add_paragraph(row[2])

        else:
            print('Нет результата. В базе данных пусто: проверьте введенные значения')

        doc.save(self.episode_name + '.docx')

        # приводим в порядок форматирование в документе
        doc_read = docx.Document(self.episode_name + '.docx')
        symbols = {"  ": "\n"}
        for i in symbols:
            for p in doc_read.paragraphs:
                if p.text.find(i) >= 0:
                    p.text = p.text.replace(i, symbols[i])

        doc_read.save(self.episode_name + '.docx')

    def calculate_middle_time(self):
        """Считаем среднее время моего ответа"""
        # мое место в очереди (индекс)
        my_pos_priority_index = self.episode_priority.index(self.my_name)
        # индекс предыдущего отписавшего
        prev_index = my_pos_priority_index - 1
        # имя предыдущего отписавшего
        prev_gamer = self.episode_priority[prev_index]

        def get_prev_times(name):
            """Кидаем запросы в бд, чтобы получить время"""
            # запрос
            sql_request = '''SELECT post_date FROM posts WHERE author_name=? AND episode_name=? ORDER BY post_date'''
            # аргумент запроса
            sql_params = ([name, self.episode_name])
            # передаем в бд
            selected = update_bd(sql_request, sql_params)
            # получаем выходные данные
            rows = selected.fetchall()

            return rows

        # время отписи предыдущего игрока
        prev_times = get_prev_times(prev_gamer)
        # мое время
        my_times = get_prev_times(self.my_name)

        # нормализуем размер двух списков со временем
        if len(prev_times) > len(my_times):
            prev_times[len(my_times):len(prev_times)] = []
        elif len(prev_times) < len(my_times):
            my_times = my_times[1:len(prev_times)+1]

        # если они равны, то
        if len(prev_times) == len(my_times):
            times_prev_real = []
            for i, k in zip(prev_times, my_times):
                a = datetime.fromtimestamp(i[0])
                b = datetime.fromtimestamp(k[0])
                if b > a:
                    times_prev_real.append(b - a)
                else:
                    times_prev_real.append(a - b)
            # считаем среднее время
            middle_time = sum((c for c in times_prev_real), timedelta()) / len(times_prev_real)
            mm, ss = divmod(middle_time.seconds, 60)
            hh, mm = divmod(mm, 60)
            middle_time = 'Мое среднее время написания поста: {} д. {} час. {} мин. {} сек.'.format(middle_time.days,
                                                                                                    hh, mm, ss)
            print(middle_time)
        else:
            return False


if __name__ == "__main__":
    # """Массив методов класса, получающих данные по эпизоду с форума и пишущих в бд"""
    # start = EpChecker()
    # # создаем бд, если ее нет
    # create_bd_file()
    # # сохраняем дату последнего апдейта
    # updating_refresh_time()
    # # создаем вебдрайвер
    # start.create_driver()
    # # получаем параметры пользователя
    # start.get_user_data(login="Oono Akira", password="xXkbX01B")
    # # получаем ссылку на эпизод
    # start.get_episode_url('http://freshair.rusff.ru/viewtopic.php?id=580')
    # start.get_planned_gamers(["Aengus Gallagher", "Oono Akira"])
    # # переходим в эпизод
    # start.go_to_url()
    # # логинимся
    # start.login()
    # # получаем идентификатор профиля
    # start.get_my_profile_id()
    # # сохраняем данные юзера в бд
    # start.save_user_params()
    # # снова идем в эпизод
    # start.go_to_url()
    # # получаем имя эпизода
    # start.get_ep_name()
    # # получаем все посты из эпизода
    # start.get_all_posts_from_ep()
    # # получаем информацию об игроках и их id в эпизоде,
    # # а таже сравниваем число игроков с запланированным
    # start.get_information_ep()
    # # получаем кол-во постов и дату последнего поста
    # start.get_episode_params()


    """Массив методов проверки существующей бд"""
    check = Base_checker('Atrocity Exhibition', 'Oono Akira', ["Aengus Gallagher", "Oono Akira"], True)
    check.post_count_in_ep()
    check.last_post_date_in_ep()
    check.my_last_post_date_in_ep()
    check.check_strict_priority()
    check.get_priority()
    check.check_last_post()
    check.get_posts_from_ep()
    check.calculate_middle_time()

    # start.driver.close()

